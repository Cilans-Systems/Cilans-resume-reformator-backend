import fitz
import os
import pdfplumber
import re
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx2pdf import convert
from openai import OpenAI

from dotenv import load_dotenv

from .funtions_for_helper import (
    formate_message_tallahassee,
    formate_message_turnpike,
    formate_message_tallahassee_without_ai_generated_summary,
    formate_message_turnpike_without_ai_generated_summary,
)

# Load environment variables from .env file
load_dotenv()

# Assuming this script is located in 'utils'
base_dir = os.path.dirname(os.path.abspath(__file__))
static_dir = os.path.join(base_dir, "../static")

# Constants
FONT_NAME = "Times New Roman"
FONT_SIZE = 10
SPACING = 10
LOGO_PATH = os.path.join(static_dir, "kyralogo.png")
CONTACT_INFO = """3673 Coolidge Ct.,
Tallahassee, FL 32311
Phone: (850) 459-5854
Email: vpatel@KyraSolutions.com"""

open_api_key = os.getenv("OPENAI_API_KEY")

client = OpenAI(api_key=open_api_key)


def extract_content_from_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""
    return text


def reformat_resume(content, keywords, resume_formate, ai_generated_summary):

    if resume_formate == "Tallahassee":

        if ai_generated_summary == "True":
            return formate_message_tallahassee(content, keywords)
        else:
            return formate_message_tallahassee_without_ai_generated_summary(
                content, keywords
            )

    else:
        if ai_generated_summary == "True":
            return formate_message_turnpike(content, keywords)
        else:
            return formate_message_turnpike_without_ai_generated_summary(
                content, keywords
            )


# /////////////////////////// DOCX to HTML Tallahassee ///////////////////////////


def add_header_with_logo_and_contact(doc):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False

        # Create table with 1 row and 2 columns
        table = header.add_table(1, 2, width=Cm(20.32))
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        table.columns[0].width = Cm(7.62)  # Adjust column width for the logo
        table.columns[1].width = Cm(12.7)  # Adjust column width for the contact info

        # Left cell for logo
        left_cell = table.cell(0, 0)
        left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if os.path.exists(LOGO_PATH):
            paragraph = left_cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(
                LOGO_PATH, width=Cm(4.67), height=Cm(2.3)
            )  # Adjust logo size
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            print(f"Logo file not found at {LOGO_PATH}")

        # Right cell for contact info
        right_cell = table.cell(0, 1)
        right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        contact_paragraph = right_cell.paragraphs[0]
        contact_run = contact_paragraph.add_run(CONTACT_INFO)
        contact_run.font.size = Pt(FONT_SIZE)
        contact_run.font.name = FONT_NAME
        contact_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT
        )  # Align text to the right

    # Add a line break after the logo
    doc.add_paragraph()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.left_margin = Cm(1.27)
        section.right_margin = Cm(1.27)

    # Remove extra space before/after paragraphs
    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)


def add_paragraph_tallahassee(
    doc, text, style=None, bold=False, alignment=None, color=None
):
    p = doc.add_paragraph(text, style=style)
    if alignment:
        p.alignment = alignment
    run = p.runs[0]
    run.bold = bold
    run.font.color.rgb = color if color else RGBColor(0, 0, 0)  # Set text color
    run.font.name = FONT_NAME  # Set font to Times New Roman
    run.font.size = Pt(FONT_SIZE)  # Set font size to 10


def add_list_item_tallahassee(doc, element, indent):
    p = doc.add_paragraph(style="List Bullet")
    for child in element.children:
        handle_element_tallahassee(doc, child, p)
    p.paragraph_format.left_indent = Cm(indent)  # Adjust the indent as needed


def handle_element_tallahassee(doc, element, parent_paragraph=None):

    if isinstance(element, str):
        if parent_paragraph:
            run = parent_paragraph.add_run(element)
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)
            run.font.color.rgb = RGBColor(0, 0, 0)
        return

    # Handle <b> and <strong> tags for bold text
    if element.name in ["b", "strong"]:
        if parent_paragraph:
            run = parent_paragraph.add_run(element.get_text())
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)  # Set font size to 10
            run.font.color.rgb = RGBColor(0, 0, 0)
            if element.name == "strong":
                parent_paragraph.paragraph_format.left_indent = Cm(0.64)
                parent_paragraph.paragraph_format.space_before = Pt(SPACING)
                parent_paragraph.paragraph_format.space_after = Pt(SPACING)
        else:
            add_paragraph_tallahassee(doc, element.get_text(), bold=True)
        return

    if element.name == "h1":
        add_paragraph_tallahassee(
            doc, element.get_text(), bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
    elif element.name == "role_title":
        add_paragraph_tallahassee(
            doc,
            element.get_text(),
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            color=RGBColor(0, 0, 255),
        )
    elif element.name == "h2":
        p = doc.add_paragraph(element.get_text())
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_format = p.paragraph_format
        p_format.space_before = Pt(SPACING)  # Add space before h2
        p_format.space_after = Pt(SPACING)  # Add space after h2
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
        run.font.name = FONT_NAME  # Set font to Times New Roman
        run.font.size = Pt(FONT_SIZE)  # Set font size to 10
    elif element.name == "h3":
        add_paragraph_tallahassee(doc, element.get_text(), style="Heading 2")
    elif element.name == "h4":
        add_paragraph_tallahassee(doc, element.get_text(), style="Heading 3")
    elif element.name == "p":
        p = doc.add_paragraph()
        for child in element.children:
            handle_element_tallahassee(doc, child, p)
    elif element.name == "table":
        # Create table with a specific number of columns (based on <th> tags)
        table = doc.add_table(
            rows=1, cols=len(element.find_all("th")), style="Table Grid"
        )

        # Disable autofit to allow manual control over the column widths
        table.autofit = False

        # Set the column widths (e.g., 1 inch for each column, adjust as needed)
        for column in table.columns:
            for cell in column.cells:
                cell.width = Cm(2.54)  # Adjust column width as needed

        # Set header row with custom font and size
        hdr_cells = table.rows[0].cells
        for idx, th in enumerate(element.find_all("th")):
            hdr_cells[idx].text = th.get_text()
            for paragraph in hdr_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(FONT_SIZE)  # Set font size to 10pt
                    run.font.name = FONT_NAME  # Set font to Times New Roman

        # Add table rows and handle table data
        for tr in element.find_all("tr")[1:]:
            row_cells = table.add_row().cells
            for idx, td in enumerate(tr.find_all("td")):
                cell_paragraph = row_cells[idx].paragraphs[0]
                handle_element_tallahassee(doc, td, cell_paragraph)

                # Align paragraph and vertically center content in each cell
                row_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Make the first column bold
        for row in table.rows:
            first_col_cell = row.cells[0]
            for paragraph in first_col_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Apply indentation to the table (e.g., 1 inch indent)
        tbl = table._tbl  # Access the underlying table element
        tblPr = tbl.tblPr  # Get the table properties
        tblInd = OxmlElement("w:tblInd")  # Create table indentation element
        tblInd.set(qn("w:w"), "500")  # Set indentation value in twips (1440 = 1 inch)
        tblInd.set(qn("w:type"), "dxa")  # Set measurement type to dxa (twips)

        # Append the indentation element to the table properties
        tblPr.append(tblInd)
    elif element.name in ["ul", "ol"]:
        indent = 1.27
        if element.name == "ol":
            indent = 2.12
        for li in element.find_all("li"):
            add_list_item_tallahassee(doc, li, indent)
        # Add space after the list
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(SPACING)
    elif element.name == "br":
        doc.add_paragraph()
    else:
        for child in element.children:
            handle_element_tallahassee(doc, child, parent_paragraph)


def convert_html_to_docx_tallahassee(html_content):
    soup = BeautifulSoup(html_content, "html.parser")
    doc = Document()

    add_header_with_logo_and_contact(doc)

    for element in soup.body:
        handle_element_tallahassee(doc, element)

    return doc


# //////////////////////////////////////////////////////////////////////

# /////////////////////////// DOCX to HTML turnpike ///////////////////////////


def convert_html_to_docx_turnpike(html_content):

    soup = BeautifulSoup(html_content, "html.parser")

    doc = Document()

    def add_paragraph(text, style=None, bold=False, alignment=None):
        p = doc.add_paragraph(text, style=style)
        if alignment:
            p.alignment = alignment
        run = p.runs[0]
        run.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
        run.font.name = FONT_NAME  # Set font to Times New Roman

    def handle_bold_tags(element, parent_paragraph):
        """Handle <b> and <strong> tags to make text bold."""
        if element.name in ["b", "strong"]:
            run = parent_paragraph.add_run(element.get_text())
            run.bold = True
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)
            run.font.color.rgb = RGBColor(0, 0, 0)
        else:
            run = parent_paragraph.add_run(element.get_text())
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE)
            run.font.color.rgb = RGBColor(0, 0, 0)

    def handle_element(element, parent_paragraph=None):
        if isinstance(element, str):
            if parent_paragraph:
                run = parent_paragraph.add_run(element)
                run.font.name = FONT_NAME
                run.font.size = Pt(FONT_SIZE)
                run.font.color.rgb = RGBColor(0, 0, 0)
            return

        if element.name == "h1":
            add_paragraph(
                "Kyra Solutions, Inc.",
                bold=True,
                alignment=WD_ALIGN_PARAGRAPH.LEFT,
            )
            add_paragraph(
                element.get_text(), bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT
            )
        elif element.name == "h2":
            add_paragraph(
                element.get_text(), bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT
            )
        elif element.name == "p":
            p = doc.add_paragraph()
            for child in element.children:
                if child.name == "strong":
                    run = p.add_run(child.get_text())
                    run.bold = True
                    run.font.name = FONT_NAME
                    run.font.size = Pt(FONT_SIZE)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    p.paragraph_format.left_indent = Cm(0.64)
                    p.paragraph_format.space_before = Pt(SPACING)
                    p.paragraph_format.space_after = Pt(SPACING)
                elif child.name == "contact":
                    # Add the contact information after the title at the right side
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run(element.get_text())
                    run.bold = True
                    run.font.name = FONT_NAME
                    run.font.size = Pt(FONT_SIZE)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                elif child.name == "role_title":
                    # Add the contact information after the title at the right side
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run = p.add_run(element.get_text())
                    run.bold = True
                    run.font.name = FONT_NAME
                    run.font.size = Pt(FONT_SIZE)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                else:
                    handle_element(child, p)
        elif element.name in ["strong", "b"]:
            if parent_paragraph:
                handle_bold_tags(element, parent_paragraph)
            else:
                add_paragraph(element.get_text(), bold=True)
        elif element.name == "table":
            table = doc.add_table(
                rows=1, cols=len(element.find_all("th")), style="Table Grid"
            )
            table.autofit = True
            hdr_cells = table.rows[0].cells
            for idx, th in enumerate(element.find_all("th")):
                hdr_cells[idx].text = th.get_text()

            for tr in element.find_all("tr")[1:]:
                row_cells = table.add_row().cells
                for idx, td in enumerate(tr.find_all("td")):
                    p = row_cells[idx].paragraphs[0]
                    for child in td.children:
                        handle_element(child, p)
                    row_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    row_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Make the first column bold
            for row in table.rows:
                first_col_cell = row.cells[0]
                for paragraph in first_col_cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Apply indentation to the table (e.g., 1 inch indent)
            tbl = table._tbl  # Access the underlying table element
            tblPr = tbl.tblPr  # Get the table properties
            tblInd = OxmlElement("w:tblInd")  # Create table indentation element
            tblInd.set(
                qn("w:w"), "500"
            )  # Set indentation value in twips (1440 = 1 inch)
            tblInd.set(qn("w:type"), "dxa")  # Set measurement type to dxa (twips)

            # Append the indentation element to the table properties
            tblPr.append(tblInd)

            # Add space after the table
            p = doc.add_paragraph()

        elif element.name in ["ul", "ol"]:
            indent = 1.27
            if element.name == "ol":
                indent = 2.12
            for li in element.find_all("li"):
                p = doc.add_paragraph(style="ListBullet")
                for child in li.children:
                    handle_element(child, p)
                p.paragraph_format.left_indent = Cm(indent)
            # Add space after the list
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(SPACING)
        elif element.name == "br":
            doc.add_paragraph()

    for element in soup.body:
        handle_element(element)

    return doc


# //////////////////////////////////////////////////////////////////////


def read_pdf(file_path):
    """Reads a .pdf file and returns its content as a string.

    Args:
        file_path (str): The path to the PDF file.

    Returns:
        str: The content of the PDF file as a string.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"The file {file_path} does not exist.")

    try:
        # Open the PDF file
        doc = fitz.open(file_path)
        content = []

        # Iterate through each page
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            content.append(page.get_text())

        return "\n".join(content)

    except Exception as e:
        raise RuntimeError(f"An error occurred while reading the PDF file: {e}")
